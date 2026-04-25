import streamlit as st
import time
import io
import tempfile
import os
import traceback
import shutil
from docx import Document
from pypdf import PdfReader, PdfWriter

# Try multiple import paths so the app works with different GenAI client packages
try:
    import genai
except Exception:
    try:
        import google.genai as genai
    except Exception:
        try:
            import google.generativeai as genai
        except Exception:
            genai = None

# --- Prompt Constants ---
AUDIO_PROMPT = """Role: You are an expert Jain Literature Archivist and Professional Transcriber.

Task: Listen to the attached audio and provide a verbatim (word-for-word) Hindi transcription with specific attention to Jain terminology.

Instructions:
1. Mantra Accuracy: Follow standard Jain Prakrit spellings for all mantras. For example, use 'णमो अरिहंताणं' instead of 'नमो अरिहंतों', and 'णमो लोए सव्व साहूणं'.
2. Terminology: Ensure correct spelling for words like 'विषापहार', 'अतिशय', 'जिन शासन', 'तीर्थंकर', 'समवशरण', 'अरिहंत', 'सिद्ध', 'आचार्य', 'उपाध्याय', and 'साधु'.
3. Purity of Language: Preserve the original mix of Sanskrit, Prakrit, and Hindi. Do not translate; only transcribe.

Expected Output: A clean, well-structured document in Unicode Hindi (Devanagari)."""

PDF_PROMPT = """Role: You are an expert Hindi archivist.

Task: Transcribe every word from the attached PDF with 100% accuracy.

Instructions:
1. Capture exact Devanagari characters including honorifics like 'ब्र०'.
2. Do not summarize; provide a verbatim transcription.
3. Maintain page structure by labeling each page clearly.
4. Preserve special symbols like 'ॐ' and '卐'.
5. REPETITION CHECK & PATTERN BREAKING: If you find yourself repeating the same word, phrase, or numbering sequence more than 3 times, STOP. This is a sign of a hallucination loop. Re-scan the page specifically for changes in text. Do not assume a list follows a uniform pattern; verify every single character against the visual source.
6. MULTI-COLUMN LAYOUT HANDLING: If a page contains multiple columns (like an index or bibliography), transcribe them column-by-column or row-by-row in a clear, structured list. Do not read across the columns as if they are a single sentence. If an entry is split by a page break, merge it into a single coherent line in your transcription to maintain context.
7. DENSE DATA FORMATTING: For pages containing dense indices (like 'ग्रन्थाक्रम'), prioritize a Numbered List format. Use the format [Number]. [Title] ([Page/Reference]). If the text becomes too dense to follow in a standard paragraph, force a new line for every new entry to prevent the model's logic from "stacking" words and looping."""


# --- Helper: Split PDF into chunks ---
def split_pdf(pdf_bytes: bytes, chunk_size: int, start_page: int = 0, end_page: int = None) -> list:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    total = len(reader.pages)
    end_page = end_page if end_page is not None else total
    end_page = min(end_page, total)
    chunks = []
    for s in range(start_page, end_page, chunk_size):
        writer = PdfWriter()
        for p in range(s, min(s + chunk_size, end_page)):
            writer.add_page(reader.pages[p])
        buf = io.BytesIO()
        writer.write(buf)
        chunks.append(buf.getvalue())
    return chunks


# --- Helper: Write to error log ---
def log_error(temp_dir: str, chunk_idx: int, error_msg: str, full_trace: str):
    if not temp_dir:
        return
    log_path = os.path.join(temp_dir, "error_log.txt")
    with open(log_path, 'a', encoding='utf-8') as f:
        f.write(f"\n{'='*60}\n")
        f.write(f"समय (Time): {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"भाग (Chunk): {chunk_idx + 1}\n")
        f.write(f"त्रुटि (Error): {error_msg}\n")
        f.write(f"Full Traceback:\n{full_trace}\n")


# --- UI Setup ---
st.set_page_config(page_title="Jain Pathshala AI", page_icon="🙏", layout="wide")

st.title("🙏 Digamber जैन धर्मग्रंथ एवं प्रवचन AI")
st.markdown("इस ऐप के माध्यम से आप **PDF (कैलाश यात्रा आदि)** और **MP3 (प्रवचन)** का सटीक शुद्ध हिंदी एवं प्राकृत अनुवाद कर सकते हैं।")

# Sidebar for Configuration
with st.sidebar:
    st.header("Settings")
    api_key = st.text_input("API Key दर्ज करें", type="password")
    st.info("Contact admin for API Key")
    st.divider()
    chunk_size = st.number_input(
        "PDF Chunk Size (pages per part)",
        min_value=1, max_value=100, value=15, step=1,
        help="Large PDFs will be split into parts of this many pages"
    )

# --- Logic ---
if api_key:
    if genai is None:
        st.error("Google GenAI क्लाइंट इंस्टॉल नहीं मिला। कृपया चलाएँ: `pip install google-genai` या `pip install google-generativeai` और ऐप को फिर से चालू करें।")
    else:
        try:
            client = genai.Client(api_key=api_key)
        except Exception as e:
            st.error(f"GenAI क्लाइंट बनाने में त्रुटि: {e}")
            st.stop()

        # --- Session state initialization ---
        for k, v in {
            'processing_active': False,
            'chunk_results': [],
            'chunk_count': 0,
            'chunk_bytes': [],
            'pdf_bytes': b'',
            'audio_bytes': b'',
            'current_chunk_index': 0,
            'processing_error': None,
            'processing_complete': False,
            'session_temp_dir': None,
            'chunk_temp_files': [],
            'is_pdf': False,
        }.items():
            if k not in st.session_state:
                st.session_state[k] = v

        # --- Helper: process a single chunk via Gemini ---
        def process_chunk(client, chunk_bytes, idx, total, is_pdf, temp_dir, mime_type=None):
            """Upload a chunk to Gemini, get transcription, save to disk. Returns text."""
            if is_pdf:
                upload_source = io.BytesIO(chunk_bytes)
                m_type = "application/pdf"
                prompt = PDF_PROMPT + f"\n8. CHUNK INFO: This is part {idx + 1} of {total}. Transcribe all content in this chunk."
            else:
                upload_source = io.BytesIO(chunk_bytes)
                m_type = mime_type or "audio/mpeg"
                prompt = AUDIO_PROMPT

            st.write("Gemini पर अपलोड हो रहा है...")
            sample_file = client.files.upload(file=upload_source, config={'mime_type': m_type})

            st.write("फाइल तैयार होने का इंतज़ार...")
            elapsed = 0
            while elapsed < 120:
                f_info = client.files.get(name=sample_file.name)
                state = getattr(getattr(f_info, 'state', None), 'name', None) or getattr(f_info, 'state', None)
                if str(state).upper() == 'ACTIVE':
                    break
                time.sleep(2)
                elapsed += 2

            st.write("Gemini से उत्तर प्राप्त हो रहा है...")
            max_retries = 3
            last_err = None
            response = None
            for attempt in range(1, max_retries + 1):
                try:
                    response = client.models.generate_content(
                        model='models/gemini-2.5-flash',
                        contents=[f_info, prompt],
                        config={'temperature': 0.1}
                    )
                    break
                except Exception as gen_err:
                    last_err = gen_err
                    if attempt < max_retries:
                        wait = attempt * 30
                        st.write(f"प्रयास {attempt} विफल ({gen_err}), {wait}s बाद पुनः प्रयास...")
                        time.sleep(wait)
            if response is None:
                raise Exception(f"3 प्रयासों के बाद भी विफल: {last_err}")

            if hasattr(response, 'text'):
                text = response.text
            elif hasattr(response, 'candidates'):
                try:
                    text = response.candidates[0].content
                except Exception:
                    text = str(response)
            else:
                text = str(response)

            if temp_dir:
                txt_path = os.path.join(temp_dir, f"chunk_{idx:03d}.txt")
                with open(txt_path, 'w', encoding='utf-8') as tf:
                    tf.write(text)

            try:
                client.files.delete(name=sample_file.name)
            except Exception:
                pass

            return text

        tab1, tab2, tab3 = st.tabs(["📤 नई फाइल अपलोड करें", "📂 फोल्डर से पुनः शुरू करें", "🎧 ऑडियो / YouTube"])

        # ====== TAB 1: Upload New ======
        with tab1:
            uploaded_file = st.file_uploader("अपनी फाइल अपलोड करें (PDF या MP3)", type=['pdf', 'mp3'])

            if uploaded_file:
                file_type = "PDF" if uploaded_file.name.endswith(".pdf") else "Audio"
                st.success(f"{file_type} फाइल तैयार है: {uploaded_file.name}")

                if st.button(f"प्रोसेस शुरू करें ({file_type})"):
                    # Cleanup previous temp files
                    for p in st.session_state.get('chunk_temp_files', []):
                        try:
                            os.unlink(p)
                        except OSError:
                            pass
                    td = st.session_state.get('session_temp_dir')
                    if td and os.path.isdir(td):
                        try:
                            os.rmdir(td)
                        except OSError:
                            pass

                    st.session_state.update({
                        'processing_active': True,
                        'chunk_results': [],
                        'current_chunk_index': 0,
                        'processing_error': None,
                        'processing_complete': False,
                        'chunk_temp_files': [],
                    })

                    if file_type == "PDF":
                        pdf_bytes = uploaded_file.read()
                        st.session_state['pdf_bytes'] = pdf_bytes
                        st.session_state['is_pdf'] = True
                        st.session_state['session_temp_dir'] = tempfile.mkdtemp(prefix='jain_')
                        chunks = split_pdf(pdf_bytes, chunk_size)
                        st.session_state['chunk_bytes'] = chunks
                        st.session_state['chunk_count'] = len(chunks)
                        st.session_state['chunk_results'] = [None] * len(chunks)
                        # Save PDF chunks to disk
                        for ci, cb in enumerate(chunks):
                            pdf_path = os.path.join(st.session_state['session_temp_dir'], f"chunk_{ci:03d}.pdf")
                            with open(pdf_path, 'wb') as pf:
                                pf.write(cb)
                    else:
                        st.session_state['audio_bytes'] = uploaded_file.read()
                        st.session_state['is_pdf'] = False
                        st.session_state['chunk_count'] = 1
                        st.session_state['chunk_results'] = [None]

                    st.rerun()

            # Show folder path if processing
            if st.session_state.get('session_temp_dir'):
                st.info(f"फाइलें यहाँ सहेजी गई हैं: {st.session_state['session_temp_dir']}")
                if st.button("📂 इस फोल्डर को 'फोल्डर से पुनः शुरू करें' टैब में खोलें", key="tab1_to_tab2"):
                    st.session_state['tab2_folder_path'] = st.session_state['session_temp_dir']
                    st.rerun()

            # --- Incremental results display ---
            done = [r for r in st.session_state.chunk_results if r is not None]
            if done:
                st.subheader("निकाल गया टेक्स्ट (Result):")
                for i, result in enumerate(st.session_state.chunk_results):
                    if result is not None:
                        label = f"भाग {i + 1}" if st.session_state.chunk_count > 1 else "परिणाम"
                        with st.expander(label, expanded=(i == len(done) - 1)):
                            st.text_area("", result, height=300,
                                         key=f"chunk_out_{i}", label_visibility="collapsed")
                            dl_col1, dl_col2 = st.columns(2)
                            with dl_col1:
                                # Download original chunk
                                if st.session_state.is_pdf and i < len(st.session_state.get('chunk_bytes', [])):
                                    st.download_button(
                                        label=f"📥 मूल PDF भाग {i + 1}",
                                        data=st.session_state['chunk_bytes'][i],
                                        file_name=f"chunk_{i:03d}.pdf",
                                        mime="application/pdf",
                                        key=f"dl_orig_{i}"
                                    )
                                elif not st.session_state.is_pdf and st.session_state.get('audio_bytes'):
                                    st.download_button(
                                        label=f"📥 मूल Audio",
                                        data=st.session_state['audio_bytes'],
                                        file_name="audio_original.mp3",
                                        mime="audio/mpeg",
                                        key=f"dl_orig_{i}"
                                    )
                            with dl_col2:
                                st.download_button(
                                    label=f"📄 Output भाग {i + 1}",
                                    data=result,
                                    file_name=f"chunk_{i:03d}_output.txt",
                                    mime="text/plain",
                                    key=f"dl_out_{i}"
                                )

            # --- Error display with resume + partial download ---
            if st.session_state.processing_error:
                failed_idx = st.session_state.current_chunk_index
                st.error(f"भाग {failed_idx + 1} पर त्रुटि: {st.session_state.processing_error}")

                col_resume, col_skip, col_dl = st.columns(3)

                with col_resume:
                    if st.button(f"भाग {failed_idx + 1} से पुनः शुरू करें"):
                        st.session_state['processing_error'] = None
                        st.session_state['processing_active'] = True
                        st.rerun()

                with col_skip:
                    if st.button(f"भाग {failed_idx + 1} छोड़ें, आगे बढ़ें"):
                        st.session_state['processing_error'] = None
                        st.session_state['chunk_results'][failed_idx] = "[यह भाग छोड़ा गया (skipped)]"
                        st.session_state['current_chunk_index'] = failed_idx + 1
                        if failed_idx + 1 >= st.session_state.chunk_count:
                            st.session_state['processing_complete'] = True
                        else:
                            st.session_state['processing_active'] = True
                        st.rerun()

                with col_dl:
                    partial_results = [r for r in st.session_state.chunk_results if r is not None]
                    if partial_results:
                        partial_text = "\n\n---\n\n".join(partial_results)
                        st.download_button(
                            label="📄 अब तक का परिणाम डाउनलोड करें",
                            data=partial_text,
                            file_name=f"Jain_Partial_{int(time.time())}.txt",
                            mime="text/plain"
                        )

                log_path = os.path.join(st.session_state['session_temp_dir'], "error_log.txt") \
                    if st.session_state['session_temp_dir'] else None
                if log_path and os.path.exists(log_path):
                    with open(log_path, 'r', encoding='utf-8') as lf:
                        log_content = lf.read()
                    st.download_button(
                        label="📋 Error Log डाउनलोड करें",
                        data=log_content,
                        file_name=f"Jain_ErrorLog_{int(time.time())}.txt",
                        mime="text/plain"
                    )

            # --- Processing loop: one chunk per rerun ---
            if st.session_state.processing_active:
                idx = st.session_state.current_chunk_index
                total = st.session_state.chunk_count

                st.progress(idx / total if total > 0 else 0,
                            text=f"प्रोसेसिंग: {idx}/{total} भाग पूर्ण")

                if idx < total:
                    with st.status(f"भाग {idx + 1}/{total} प्रोसेस हो रहा है...", expanded=True):
                        try:
                            chunk_data = st.session_state['chunk_bytes'][idx] if st.session_state.is_pdf else st.session_state['audio_bytes']
                            text = process_chunk(client, chunk_data, idx, total,
                                                 st.session_state.is_pdf, st.session_state['session_temp_dir'])

                            st.session_state['chunk_results'][idx] = text
                            st.session_state['current_chunk_index'] = idx + 1

                        except Exception as e:
                            full_trace = traceback.format_exc()
                            log_error(st.session_state['session_temp_dir'], idx, str(e), full_trace)
                            st.session_state['processing_error'] = str(e)
                            st.session_state['processing_active'] = False
                            st.rerun()

                    if st.session_state['current_chunk_index'] >= total:
                        st.session_state['processing_active'] = False
                        st.session_state['processing_complete'] = True

                    st.rerun()

            # --- Final combined download ---
            if st.session_state.processing_complete:
                full_text = "\n\n---\n\n".join(
                    r for r in st.session_state.chunk_results if r is not None
                )

                if st.session_state.chunk_count > 1:
                    st.subheader("पूर्ण परिणाम (Complete Result):")
                    st.text_area("Final Transcript->", full_text, height=500)

                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        label="📄 Text (.txt) डाउनलोड करें",
                        data=full_text,
                        file_name=f"Jain_Output_{int(time.time())}.txt",
                        mime="text/plain"
                    )
                with col2:
                    doc = Document()
                    for r in st.session_state.chunk_results:
                        if r:
                            doc.add_paragraph(r)
                            doc.add_paragraph()
                    docx_buffer = io.BytesIO()
                    doc.save(docx_buffer)
                    docx_buffer.seek(0)
                    st.download_button(
                        label="📋 Word (.docx) डाउनलोड करें",
                        data=docx_buffer.getvalue(),
                        file_name=f"Jain_Output_{int(time.time())}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

        # ====== TAB 2: Resume from Folder ======
        with tab2:
            folder_path = st.text_input("फोल्डर पथ दर्ज करें (Folder Path)",
                                        value=st.session_state.get('tab2_folder_path', ''),
                                        placeholder="C:\\Users\\...\\jain_XXXX")

            if folder_path and os.path.isdir(folder_path):
                # Scan for chunk source files (PDF or audio) and TXTs
                AUDIO_EXTS = ('.mp3', '.m4a', '.wav', '.ogg')
                source_files = sorted([f for f in os.listdir(folder_path)
                                       if f.startswith('chunk_') and (f.endswith('.pdf') or f.endswith(AUDIO_EXTS))])
                txt_files = {f.replace('.txt', ''): f for f in os.listdir(folder_path)
                             if f.startswith('chunk_') and f.endswith('.txt')}

                if not source_files:
                    st.warning("इस फोल्डर में कोई chunk (PDF/Audio) नहीं मिला।")
                else:
                    # Detect type from first file
                    first_ext = os.path.splitext(source_files[0])[1].lower()
                    is_pdf_folder = first_ext == '.pdf'
                    folder_type = "PDF" if is_pdf_folder else "Audio"
                    st.success(f"{len(source_files)} {folder_type} chunks मिले, {len(txt_files)} पहले से प्रोसेस हो चुके हैं।")

                    # Initialize regen state
                    if 'regen_active' not in st.session_state:
                        st.session_state['regen_active'] = False
                    if 'regen_index' not in st.session_state:
                        st.session_state['regen_index'] = -1

                    for src_file in source_files:
                        stem, ext = os.path.splitext(src_file)
                        chunk_stem = stem  # e.g. chunk_000
                        idx = int(chunk_stem.split('_')[1])
                        has_output = chunk_stem in txt_files

                        col_name, col_status, col_action, col_dl, col_view = st.columns([3, 2, 1, 2, 2])

                        with col_name:
                            st.write(f"**{chunk_stem}** ({ext})")

                        with col_status:
                            if has_output:
                                st.write(":white_check_mark: प्रोसेस हो चुका")
                            else:
                                st.write(":x: बाकी है")

                        with col_action:
                            if st.button("🔄", key=f"regen_{idx}", help=f"भाग {idx + 1} पुनः प्रोसेस करें"):
                                st.session_state['regen_active'] = True
                                st.session_state['regen_index'] = idx
                                st.session_state['regen_folder'] = folder_path
                                st.session_state['regen_is_pdf'] = is_pdf_folder
                                st.session_state['regen_ext'] = ext
                                st.rerun()

                        with col_dl:
                            # Download original chunk
                            src_path = os.path.join(folder_path, src_file)
                            with open(src_path, 'rb') as sf:
                                src_data = sf.read()
                            mime_map = {'.pdf': 'application/pdf', '.mp3': 'audio/mpeg',
                                        '.m4a': 'audio/mp4', '.wav': 'audio/wav', '.ogg': 'audio/ogg'}
                            st.download_button(
                                label=f"📥 मूल",
                                data=src_data,
                                file_name=src_file,
                                mime=mime_map.get(ext, 'application/octet-stream'),
                                key=f"dl_src_{idx}"
                            )

                        with col_view:
                            if has_output:
                                txt_path = os.path.join(folder_path, txt_files[chunk_stem])
                                with open(txt_path, 'r', encoding='utf-8') as rf:
                                    content = rf.read()
                                with st.popover("👁 देखें"):
                                    st.text_area("", content, height=300, key=f"view_{idx}", label_visibility="collapsed")
                                st.download_button(
                                    label=f"📄 Output",
                                    data=content,
                                    file_name=f"{chunk_stem}_output.txt",
                                    mime="text/plain",
                                    key=f"dl_out_t2_{idx}"
                                )

                    st.divider()

                    # Download all completed outputs combined
                    all_texts = []
                    for src_file in source_files:
                        chunk_stem = os.path.splitext(src_file)[0]
                        if chunk_stem in txt_files:
                            txt_path = os.path.join(folder_path, txt_files[chunk_stem])
                            with open(txt_path, 'r', encoding='utf-8') as rf:
                                all_texts.append(rf.read())

                    if all_texts:
                        combined = "\n\n---\n\n".join(all_texts)
                        col1, col2 = st.columns(2)
                        with col1:
                            st.download_button(
                                label="📄 सभी परिणाम Text (.txt) डाउनलोड करें",
                                data=combined,
                                file_name=f"Jain_Combined_{int(time.time())}.txt",
                                mime="text/plain",
                                key="tab2_dl_txt"
                            )
                        with col2:
                            doc = Document()
                            for t in all_texts:
                                doc.add_paragraph(t)
                                doc.add_paragraph()
                            docx_buffer = io.BytesIO()
                            doc.save(docx_buffer)
                            docx_buffer.seek(0)
                            st.download_button(
                                label="📋 सभी परिणाम Word (.docx) डाउनलोड करें",
                                data=docx_buffer.getvalue(),
                                file_name=f"Jain_Combined_{int(time.time())}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="tab2_dl_docx"
                            )

                    # --- Regenerate processing ---
                    if st.session_state.get('regen_active'):
                        regen_idx = st.session_state['regen_index']
                        regen_folder = st.session_state['regen_folder']
                        regen_is_pdf = st.session_state.get('regen_is_pdf', True)
                        regen_ext = st.session_state.get('regen_ext', '.pdf')
                        src_path = os.path.join(regen_folder, f"chunk_{regen_idx:03d}{regen_ext}")

                        mime_map = {'.pdf': 'application/pdf', '.mp3': 'audio/mpeg',
                                    '.m4a': 'audio/mp4', '.wav': 'audio/wav', '.ogg': 'audio/ogg'}

                        with st.status(f"भाग {regen_idx + 1} पुनः प्रोसेस हो रहा है...", expanded=True):
                            try:
                                with open(src_path, 'rb') as pf:
                                    chunk_data = pf.read()
                                text = process_chunk(client, chunk_data, regen_idx, len(source_files),
                                                     regen_is_pdf, regen_folder,
                                                     mime_type=None if regen_is_pdf else mime_map.get(regen_ext, 'audio/mpeg'))
                                st.session_state['regen_active'] = False
                                st.session_state['regen_index'] = -1
                                st.success(f"भाग {regen_idx + 1} सफलतापूर्वक पुनः प्रोसेस हो गया!")
                                st.rerun()
                            except Exception as e:
                                st.session_state['regen_active'] = False
                                st.error(f"त्रुटि: {e}")

            elif folder_path:
                st.error("यह फोल्डर मौजूद नहीं है। कृपया सही पथ दर्ज करें।")

        # ====== TAB 3: Audio / YouTube ======
        with tab3:
            st.markdown("""### YouTube/ऑडियो फाइल कैसे प्राप्त करें?

**विकल्प 1:** यदि आपके पास पहले से ऑडियो फाइल (MP3/M4A) है, तो सीधे नीचे अपलोड करें।

**विकल्प 2:** YouTube से ऑडियो डाउनलोड करने के लिए:
1. [y2mate.guru](https://www.y2mate.guru/) या [ssyoutube.com](https://ssyoutube.com/) पर जाएं
2. YouTube वीडियो का URL पेस्ट करें
3. **MP3/Audio** फॉर्मेट चुनें और डाउनलोड करें
4. डाउनलोड की गई फाइल नीचे अपलोड करें

**विकल्प 3:** ब्राउज़र एक्सटेंशन जैसे **Video DownloadHelper** (Firefox/Chrome) से भी डाउनलोड कर सकते हैं।
""")
            st.divider()

            # Check for ffmpeg availability
            ffmpeg_available = shutil.which("ffmpeg") is not None

            if not ffmpeg_available:
                st.warning("ffmpeg इंस्टॉल नहीं मिला। लंबी ऑडियो फाइलें बिना chunking के प्रोसेस होंगी।")

            audio_file = st.file_uploader("ऑडियो फाइल अपलोड करें (MP3, M4A, WAV, OGG)", type=['mp3', 'm4a', 'wav', 'ogg'], key="yt_audio_upload")

            if audio_file:
                st.success(f"ऑडियो फाइल तैयार है: {audio_file.name}")

                # Initialize YouTube/audio session state
                for k, v in {
                    'yt_processing_active': False,
                    'yt_chunk_results': [],
                    'yt_chunk_count': 0,
                    'yt_current_chunk_index': 0,
                    'yt_processing_error': None,
                    'yt_processing_complete': False,
                    'yt_temp_dir': None,
                    'yt_chunk_files': [],
                }.items():
                    if k not in st.session_state:
                        st.session_state[k] = v

                if st.button("प्रोसेस शुरू करें (Audio)"):
                    yt_temp = tempfile.mkdtemp(prefix='jain_yt_')
                    st.session_state['yt_temp_dir'] = yt_temp

                    audio_bytes = audio_file.read()

                    # Determine mime type from extension
                    ext = audio_file.name.rsplit('.', 1)[-1].lower()
                    mime_map = {'mp3': 'audio/mpeg', 'm4a': 'audio/mp4', 'wav': 'audio/wav', 'ogg': 'audio/ogg'}
                    audio_mime = mime_map.get(ext, 'audio/mpeg')

                    if ffmpeg_available:
                        with st.status("ऑडियो को 30 मिनट के भागों में विभाजित किया जा रहा है...", expanded=True):
                            try:
                                from pydub import AudioSegment

                                # Save uploaded file to disk for pydub
                                src_path = os.path.join(yt_temp, f"source.{ext}")
                                with open(src_path, 'wb') as sf:
                                    sf.write(audio_bytes)

                                audio_seg = AudioSegment.from_file(src_path)
                                chunk_duration_ms = 30 * 60 * 1000  # 30 minutes
                                chunks = []
                                chunk_bytes_list = []

                                for i in range(0, len(audio_seg), chunk_duration_ms):
                                    chunk = audio_seg[i:i + chunk_duration_ms]
                                    chunk_path = os.path.join(yt_temp, f"chunk_{len(chunks):03d}.{ext}")
                                    chunk.export(chunk_path, format="ipod" if ext == "m4a" else ext)
                                    with open(chunk_path, 'rb') as cf:
                                        chunk_bytes_list.append(cf.read())
                                    chunks.append(chunk_path)

                                st.write(f"{len(chunks)} भाग बनाए गए")

                                st.session_state.update({
                                    'yt_processing_active': True,
                                    'yt_chunk_results': [None] * len(chunks),
                                    'yt_chunk_count': len(chunks),
                                    'yt_current_chunk_index': 0,
                                    'yt_processing_error': None,
                                    'yt_processing_complete': False,
                                    'yt_chunk_files': chunk_bytes_list,
                                    'yt_audio_mime': audio_mime,
                                })
                                st.rerun()

                            except Exception as e:
                                st.error(f"ऑडियो chunking त्रुटि: {e}")
                    else:
                        # No ffmpeg — process as single chunk
                        st.session_state.update({
                            'yt_processing_active': True,
                            'yt_chunk_results': [None],
                            'yt_chunk_count': 1,
                            'yt_current_chunk_index': 0,
                            'yt_processing_error': None,
                            'yt_processing_complete': False,
                            'yt_chunk_files': [audio_bytes],
                            'yt_audio_mime': audio_mime,
                        })
                        st.session_state['yt_temp_dir'] = yt_temp
                        st.rerun()

            # Show temp dir path
            if st.session_state.get('yt_temp_dir'):
                st.info(f"फाइलें यहाँ सहेजी गई हैं: {st.session_state['yt_temp_dir']}")
                if st.button("📂 इस फोल्डर को 'फोल्डर से पुनः शुरू करें' टैब में खोलें", key="tab3_to_tab2"):
                    st.session_state['tab2_folder_path'] = st.session_state['yt_temp_dir']
                    st.rerun()

            # --- Incremental results display ---
            yt_done = [r for r in st.session_state.get('yt_chunk_results', []) if r is not None]
            if yt_done:
                st.subheader("निकाल गया टेक्स्ट (Result):")
                for i, result in enumerate(st.session_state['yt_chunk_results']):
                    if result is not None:
                        label = f"भाग {i + 1}" if st.session_state['yt_chunk_count'] > 1 else "परिणाम"
                        with st.expander(label, expanded=(i == len(yt_done) - 1)):
                            st.text_area("", result, height=300,
                                         key=f"yt_chunk_out_{i}", label_visibility="collapsed")
                            dl_col1, dl_col2 = st.columns(2)
                            with dl_col1:
                                if i < len(st.session_state.get('yt_chunk_files', [])):
                                    st.download_button(
                                        label=f"📥 मूल Audio भाग {i + 1}",
                                        data=st.session_state['yt_chunk_files'][i],
                                        file_name=f"chunk_{i:03d}.m4a",
                                        mime=st.session_state.get('yt_audio_mime', 'audio/mpeg'),
                                        key=f"yt_dl_orig_{i}"
                                    )
                            with dl_col2:
                                st.download_button(
                                    label=f"📄 Output भाग {i + 1}",
                                    data=result,
                                    file_name=f"chunk_{i:03d}_output.txt",
                                    mime="text/plain",
                                    key=f"yt_dl_out_{i}"
                                )

            # --- Error display with resume/skip ---
            if st.session_state.get('yt_processing_error'):
                failed_idx = st.session_state['yt_current_chunk_index']
                st.error(f"भाग {failed_idx + 1} पर त्रुटि: {st.session_state['yt_processing_error']}")

                col_resume, col_skip = st.columns(2)
                with col_resume:
                    if st.button(f"भाग {failed_idx + 1} से पुनः शुरू करें", key="yt_resume"):
                        st.session_state['yt_processing_error'] = None
                        st.session_state['yt_processing_active'] = True
                        st.rerun()
                with col_skip:
                    if st.button(f"भाग {failed_idx + 1} छोड़ें", key="yt_skip"):
                        st.session_state['yt_processing_error'] = None
                        st.session_state['yt_chunk_results'][failed_idx] = "[यह भाग छोड़ा गया (skipped)]"
                        st.session_state['yt_current_chunk_index'] = failed_idx + 1
                        if failed_idx + 1 >= st.session_state['yt_chunk_count']:
                            st.session_state['yt_processing_complete'] = True
                        else:
                            st.session_state['yt_processing_active'] = True
                        st.rerun()

            # --- Processing loop ---
            if st.session_state.get('yt_processing_active'):
                idx = st.session_state['yt_current_chunk_index']
                total = st.session_state['yt_chunk_count']

                st.progress(idx / total if total > 0 else 0,
                            text=f"प्रोसेसिंग: {idx}/{total} भाग पूर्ण")

                if idx < total:
                    with st.status(f"भाग {idx + 1}/{total} प्रोसेस हो रहा है...", expanded=True):
                        try:
                            chunk_data = st.session_state['yt_chunk_files'][idx]
                            text = process_chunk(client, chunk_data, idx, total,
                                                 False, st.session_state['yt_temp_dir'],
                                                 mime_type=st.session_state.get('yt_audio_mime', 'audio/mpeg'))

                            st.session_state['yt_chunk_results'][idx] = text
                            st.session_state['yt_current_chunk_index'] = idx + 1

                        except Exception as e:
                            full_trace = traceback.format_exc()
                            log_error(st.session_state['yt_temp_dir'], idx, str(e), full_trace)
                            st.session_state['yt_processing_error'] = str(e)
                            st.session_state['yt_processing_active'] = False
                            st.rerun()

                    if st.session_state['yt_current_chunk_index'] >= total:
                        st.session_state['yt_processing_active'] = False
                        st.session_state['yt_processing_complete'] = True

                    st.rerun()

            # --- Final combined download ---
            if st.session_state.get('yt_processing_complete'):
                full_text = "\n\n---\n\n".join(
                    r for r in st.session_state['yt_chunk_results'] if r is not None
                )

                st.subheader("पूर्ण परिणाम (Complete Result):")
                st.text_area("Final Transcript->", full_text, height=500, key="yt_final")

                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        label="📄 Text (.txt) डाउनलोड करें",
                        data=full_text,
                        file_name=f"Jain_YT_Output_{int(time.time())}.txt",
                        mime="text/plain",
                        key="yt_dl_txt"
                    )
                with col2:
                    doc = Document()
                    for r in st.session_state['yt_chunk_results']:
                        if r:
                            doc.add_paragraph(r)
                            doc.add_paragraph()
                    docx_buffer = io.BytesIO()
                    doc.save(docx_buffer)
                    docx_buffer.seek(0)
                    st.download_button(
                        label="📋 Word (.docx) डाउनलोड करें",
                        data=docx_buffer.getvalue(),
                        file_name=f"Jain_YT_Output_{int(time.time())}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="yt_dl_docx"
                    )

else:
    st.warning("ऐप चलाने के लिए कृपया साइडबार में API Key डालें।")
